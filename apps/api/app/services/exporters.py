from __future__ import annotations

from io import BytesIO

import pandas as pd
from docx import Document

from app.models import InquiryResult


def _items_to_frame(result: InquiryResult) -> pd.DataFrame:
    rows = []
    for item in result.items:
        rows.append(
            {
                "清单编码": item.boq_code or "",
                "清单项目": item.name,
                "专业章节": item.category or "",
                "项目特征": item.specification or "",
                "材质": item.material or "",
                "工程量": item.quantity,
                "单位": item.unit,
                "询价方式": item.inquiry_method or "",
                "来源图纸": ", ".join(item.source_documents),
                "参考供应商": item.reference_vendor or "",
                "参考单价": item.reference_unit_price or "",
                "参考合价": item.reference_total_price or "",
                "参考依据": item.reference_basis or "",
                "异常": ", ".join(item.anomalies),
                "来源摘录": item.source_snippet,
            }
        )
    return pd.DataFrame(rows)


def export_xlsx(result: InquiryResult) -> BytesIO:
    workbook = BytesIO()
    items_frame = _items_to_frame(result)
    manufacturer_df = items_frame[items_frame["询价方式"] == "厂家询价"]
    market_df = items_frame[items_frame["询价方式"] == "市场询价"]

    summary_frame = pd.DataFrame(
        [
            {
                "项目名称": result.project_name,
                "项目数": result.summary.item_count,
                "参考价项": result.summary.reference_count,
                "待询价项": result.summary.pending_count,
                "异常项": result.summary.flagged_count,
                "参考合计": result.summary.reference_subtotal,
                "币种": result.summary.currency,
                "抽取方式": result.extraction_mode,
                "定价模式": result.pricing_mode,
            }
        ]
    )

    with pd.ExcelWriter(workbook, engine="openpyxl") as writer:
        summary_frame.to_excel(writer, sheet_name="汇总", index=False)
        items_frame.to_excel(writer, sheet_name="项目清单", index=False)

        if not manufacturer_df.empty:
            manufacturer_df.to_excel(writer, sheet_name="厂家询价项", index=False)

        if not market_df.empty:
            market_df.to_excel(writer, sheet_name="市场询价项", index=False)

    workbook.seek(0)
    return workbook


def export_docx(result: InquiryResult) -> BytesIO:
    document = Document()
    document.add_heading("项目询价清单", level=0)
    document.add_paragraph(f"项目名称: {result.project_name}")
    document.add_paragraph(f"请求编号: {result.request_id}")
    document.add_paragraph(f"抽取方式: {result.extraction_mode}")
    document.add_paragraph(
        f"共 {result.summary.item_count} 项，参考价 {result.summary.reference_count} 项，"
        f"待询价 {result.summary.pending_count} 项，异常 {result.summary.flagged_count} 项，"
        f"参考合计 {result.summary.reference_subtotal:.2f} {result.summary.currency}"
    )

    document.add_heading("来源文件", level=1)
    for source in result.documents:
        document.add_paragraph(
            f"{source.filename} | {source.file_type} | {source.parser} | {source.text_excerpt}",
            style="List Bullet",
        )

    document.add_heading("项目清单", level=1)
    table = document.add_table(rows=1, cols=8)
    table.style = "Light List Accent 1"
    header_cells = table.rows[0].cells
    headers = ["清单编码", "名称", "项目特征", "工程量", "询价方式", "参考单价", "参考合价", "异常"]
    for cell, header in zip(header_cells, headers):
        cell.text = header

    for item in result.items:
        row = table.add_row().cells
        row[0].text = item.boq_code or "-"
        row[1].text = item.name
        row[2].text = item.specification or "-"
        row[3].text = f"{item.quantity:g} {item.unit}"
        row[4].text = item.inquiry_method or "-"
        row[5].text = f"{item.reference_unit_price:.2f}" if item.reference_unit_price is not None else "-"
        row[6].text = f"{item.reference_total_price:.2f}" if item.reference_total_price is not None else "-"
        row[7].text = ", ".join(item.anomalies) or "-"

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
